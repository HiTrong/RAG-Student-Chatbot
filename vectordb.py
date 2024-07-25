# ====== Import Library ======
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceInstructEmbeddings
from docx import Document
from langchain_core.documents import Document as LC_Document
from pypdf import PdfReader
import torch
import os
import yaml

# ====== Load configuration ======
with open("db_config.yml","r") as f:
    db_config = yaml.safe_load(f)
    
with open("model_config.yml","r") as f:
    model_config = yaml.safe_load(f)
    
# ====== Load Embedding model ======
def load_embedding(embedding_path=model_config["embedding_path"]):
    return HuggingFaceInstructEmbeddings(model_name=embedding_path)

# ====== Get List Documents =====
def get_list_documents():
    documents = []
    for root, dirs, files in os.walk(db_config["word_path"]):
        for file in files:
            if file.endswith(".docx"):
                documents.append(file)

    for root, dirs, files in os.walk(db_config["pdf_path"]):
        for file in files:
            if file.endswith(".pdf"):
                documents.append(file)
    return documents

# ====== Get File Documents ======
def get_document(filename):
    if filename.endswith(".docx"):
        filepath = os.path.join(db_config["word_path"], filename)
    if filename.endswith(".pdf"):
        filepath = os.path.join(db_config["pdf_path"], filename)
    with open(filepath, "rb") as file:
        return file.read()

# ====== Delete File Documents ======
def delete_document(filename):
    try:
        if filename.endswith(".docx"):
            filepath = os.path.join(db_config["word_path"], filename)
        if filename.endswith(".pdf"):
            filepath = os.path.join(db_config["pdf_path"], filename)
        os.remove(filepath)
        return True
    except:
        return False
    
# ====== Get Details Document ======
def get_details(filename):
    if filename.endswith(".docx"):
        filepath = os.path.join(db_config["word_path"], filename)
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        
    if filename.endswith(".pdf"):
        filepath = os.path.join(db_config["pdf_path"], filename)
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    
    file_details = {
        "Tên file": filename,
        "Đường dẫn": filepath,
        "Kích thước": os.path.getsize(filepath)
    }
    
    return file_details, text

# ====== My PyWordLoader ======
class MyPyWordLoader:
    def __init__(self, word_path=db_config["word_path"]):
        self.word_path = word_path
    
    def load(self):
        docx_documents = []
        for root, dirs, files in os.walk(self.word_path):
            for file in files:
                if file.endswith(".docx"):
                    file_path = os.path.join(root, file)
                    document = Document(file_path)
                    text = "\n".join([para.text for para in document.paragraphs if para.text.strip() != ""])
                    docx_documents.append(LC_Document(page_content=text))
                    
        return docx_documents

# ====== Create VectorDB (FAISS) ======
def create_vectordb_with_text(text: str, 
                              chunk_size=db_config["database_config"]["chunk_size"], 
                              separator=db_config["database_config"]["separator"], 
                              chunk_overlap=db_config["database_config"]["chunk_overlap"], 
                              db_path=db_config["database_path"]):
    text_splitter = CharacterTextSplitter(
        separator=separator,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    
    chunks = text_splitter.split_text(text=text)
    
    # Load Embedding
    embedding = load_embedding(model_config["embedding_path"])
    
    db = FAISS.from_texts(texts=chunks, embedding=embedding)
    db.save_local(db_path)
    
def create_vectordb_with_file(pdf_path=db_config["pdf_path"], 
                              word_path=db_config["word_path"], 
                              chunk_size=db_config["database_config"]["chunk_size"], 
                              chunk_overlap=db_config["database_config"]["chunk_overlap"], 
                              db_path=db_config["database_path"]):
    
    os.makedirs(pdf_path, exist_ok=True)
    os.makedirs(word_path, exist_ok=True)
    
    # pdf proccess
    pdf_loader = DirectoryLoader(pdf_path, glob="*.pdf", loader_cls = PyPDFLoader)
    pdf_documents = pdf_loader.load()

    # word proccess
    word_loader = MyPyWordLoader(word_path)
    word_documents = word_loader.load()
    
    documents = pdf_documents + word_documents
    if len(documents) == 0:
        print("No documents were found!")
        return
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = text_splitter.split_documents(documents)
    
    # # Load Embedding
    embedding = load_embedding(model_config["embedding_path"])
    
    db = FAISS.from_documents(documents=chunks, embedding=embedding)
    db.save_local(db_path)
    
# ====== Load VectorDB ======
def load_vector_db(db_path=db_config["database_path"]):
    embedding = load_embedding(model_config["embedding_path"])
    db = FAISS.load_local(db_path, embedding, allow_dangerous_deserialization=True)
    return db

    
    